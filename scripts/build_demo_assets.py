"""
Build demo assets: three .docx templates + a sample tenants.xlsx.

Re-run safely (idempotent overwrites).
"""
from pathlib import Path

from docx import Document
from openpyxl import Workbook

ROOT = Path(__file__).resolve().parent.parent
PACK = ROOT / 'docforge' / 'template_packs' / 'housing_tribunal_pack'
SAMPLES = ROOT / 'docforge' / 'samples'
SAMPLES.mkdir(parents=True, exist_ok=True)


def _write_docx(path: Path, title: str, body_lines: list[str]):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in body_lines:
        doc.add_paragraph(line)
    doc.save(str(path))


_write_docx(
    PACK / 'legal_form.docx',
    'Housing Tribunal — Legal Form',
    [
        'Tenant: {{ tenant_name }}',
        'Address: {{ address }}',
        'Postal code: {{ postal_code }}',
        'Monthly rent: {{ rent_amount }}',
        'Arrears owed: {{ arrears_amount }}',
        '',
        'This form is submitted to the housing tribunal on behalf of the above-named tenant.',
    ],
)

_write_docx(
    PACK / 'supporting_list.docx',
    'Supporting Documents List',
    [
        'Case: {{ tenant_name }} — {{ address }}',
        '',
        '1. Signed tenancy agreement',
        '2. Rent ledger showing arrears of {{ arrears_amount }}',
        '3. Copies of demand letters',
        '4. Proof of service',
    ],
)

_write_docx(
    PACK / 'bailiff_sheet.docx',
    'Bailiff Instruction Sheet',
    [
        'Subject: {{ tenant_name }}',
        'Service address: {{ address }}, {{ postal_code }}',
        '',
        'Please attend the above property and serve the attached tribunal notice.',
        'Report any non-service back to the instructing property manager.',
    ],
)


wb = Workbook()
ws = wb.active
ws.title = 'Tenants'
ws.append([
    'Tenant Name', 'Address', 'Postal Code', 'Rent Amount', 'Arrears Amount', 'Email',
])
ws.append(['Alice Johnson', '12 Maple Street', 'SW1A 1AA', '1200', '3600', 'alice@example.com'])
ws.append(['Bob Martinez', '48 Oak Avenue', '', '950', '1900', 'bob@example.com'])
ws.append(['Carol Ng', '7 Birch Lane', 'E14 5AB', '', '2100', 'carol@example.com'])
wb.save(str(SAMPLES / 'tenants.xlsx'))

print('Built:', PACK, SAMPLES / 'tenants.xlsx')
